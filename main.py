import os
import json
import logging
from datetime import datetime
from telegram import Update, ReplyKeyboardMarkup, InlineKeyboardMarkup, InlineKeyboardButton
from telegram.ext import Application, CommandHandler, MessageHandler, CallbackQueryHandler, ContextTypes, filters
from dotenv import load_dotenv
import ai_config
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO

# تنظیم لاگینگ
logging.basicConfig(
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    level=logging.INFO
)
logger = logging.getLogger(__name__)

# بارگذاری متغیرهای محیطی
load_dotenv()
TOKEN = os.getenv("TELEGRAM_TOKEN", "7595888832:AAGHkNqZcQZ4RDn5ww7vtYMPpNdiXmOpg7c")
ADMIN_USER_ID = int(os.getenv("ADMIN_USER_ID", "86551786"))

# مسیر فایل کاربران
USERS_FILE = "users.json"
LOGS_FILE = "logs.json"
REPORTS_FILE = "reports.json"

# حالت‌های مختلف کاربر
WAITING_FOR_PASSKEY = "waiting_for_passkey"
WAITING_FOR_REPORT = "waiting_for_report"
WAITING_FOR_CONFIRMATION = "waiting_for_confirmation"
NORMAL_MODE = "normal_mode"

# دکمه‌های منو
MENU_KEYBOARD = [
    ["ارسال گزارش/نامه", "راهنما"],
    ["آرشیو", "ارتباط با مدیر"]
]

# تنظیم passkey پیش‌فرض (در حالت واقعی باید از یک سیستم امن‌تر استفاده شود)
DEFAULT_PASSKEYS = {
    "admin": "admin123",
    "user1": "user123"
}

# بارگذاری کاربران از فایل
def load_users():
    try:
        with open(USERS_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    except (FileNotFoundError, json.JSONDecodeError):
        # ایجاد فایل اولیه با کاربر ادمین
        users = {
            str(ADMIN_USER_ID): {
                "role": "admin",
                "passkey": DEFAULT_PASSKEYS["admin"],
                "failed_attempts": 0,
                "is_authenticated": False
            }
        }
        save_users(users)
        return users

# ذخیره کاربران در فایل
def save_users(users):
    with open(USERS_FILE, 'w', encoding='utf-8') as f:
        json.dump(users, f, ensure_ascii=False, indent=4)

# ثبت لاگ
def log_event(user_id, event_type, details=None):
    try:
        with open(LOGS_FILE, 'r', encoding='utf-8') as f:
            logs = json.load(f)
    except (FileNotFoundError, json.JSONDecodeError):
        logs = []
    
    logs.append({
        "timestamp": datetime.now().isoformat(),
        "user_id": user_id,
        "event_type": event_type,
        "details": details
    })
    
    with open(LOGS_FILE, 'w', encoding='utf-8') as f:
        json.dump(logs, f, ensure_ascii=False, indent=4)

# ذخیره گزارش
def save_report(user_id, original_text, ai_response):
    try:
        with open(REPORTS_FILE, 'r', encoding='utf-8') as f:
            reports = json.load(f)
    except (FileNotFoundError, json.JSONDecodeError):
        reports = []
    
    report_id = len(reports) + 1
    reports.append({
        "id": report_id,
        "timestamp": datetime.now().isoformat(),
        "user_id": user_id,
        "original_text": original_text,
        "ai_response": ai_response
    })
    
    with open(REPORTS_FILE, 'w', encoding='utf-8') as f:
        json.dump(reports, f, ensure_ascii=False, indent=4)
    
    return report_id

# ایجاد فایل ورد با محتوای گزارش
def create_docx(report_content, report_id):
    doc = docx.Document()
    
    # اضافه کردن سربرگ
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header.add_run("شورای صنفی دانشگاه اصفهان")
    header_run.bold = True
    header_run.font.size = Pt(16)
    
    # اضافه کردن تاریخ
    date_paragraph = doc.add_paragraph()
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    date_paragraph.add_run(f"تاریخ: {datetime.now().strftime('%Y/%m/%d')}")
    
    # اضافه کردن شماره گزارش
    id_paragraph = doc.add_paragraph()
    id_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    id_paragraph.add_run(f"شماره گزارش: {report_id}")
    
    # اضافه کردن خط جداکننده
    doc.add_paragraph("_" * 50)
    
    # اضافه کردن محتوای گزارش
    content_paragraph = doc.add_paragraph()
    content_paragraph.add_run(report_content)
    
    # اضافه کردن امضا
    signature = doc.add_paragraph()
    signature.alignment = WD_ALIGN_PARAGRAPH.LEFT
    signature.add_run("\n\nبا احترام\nشورای صنفی دانشگاه اصفهان")
    
    # ذخیره در بایت آرایه
    f = BytesIO()
    doc.save(f)
    f.seek(0)
    return f

# شروع بات
async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = str(update.effective_user.id)
    users = load_users()
    
    # بررسی وضعیت احراز هویت کاربر
    if user_id in users and users[user_id].get("is_authenticated", False):
        # کاربر قبلاً احراز هویت شده است
        await show_menu(update, context)
    else:
        # کاربر نیاز به احراز هویت دارد
        context.user_data["state"] = WAITING_FOR_PASSKEY
        await update.message.reply_text(
            "به ربات شورای صنفی دانشگاه اصفهان خوش آمدید.\n"
            "لطفاً پس‌کی خود را وارد کنید:"
        )
    
    # ثبت لاگ شروع استفاده از بات
    log_event(user_id, "start_bot")

# نمایش منوی اصلی
async def show_menu(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data["state"] = NORMAL_MODE
    await update.message.reply_text(
        "لطفاً یکی از گزینه‌های زیر را انتخاب کنید:",
        reply_markup=ReplyKeyboardMarkup(MENU_KEYBOARD, resize_keyboard=True)
    )

# پردازش پیام‌های متنی
async def handle_message(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = str(update.effective_user.id)
    text = update.message.text
    state = context.user_data.get("state", WAITING_FOR_PASSKEY)
    users = load_users()
    
    # پردازش بر اساس حالت کاربر
    if state == WAITING_FOR_PASSKEY:
        # بررسی اعتبار پس‌کی
        is_valid = False
        for role, passkey in DEFAULT_PASSKEYS.items():
            if text == passkey:
                is_valid = True
                user_role = role
                break
        
        if is_valid:
            # پس‌کی معتبر است
            if user_id not in users:
                users[user_id] = {
                    "role": "user" if user_role != "admin" else "admin",
                    "passkey": text,
                    "failed_attempts": 0,
                    "is_authenticated": True
                }
            else:
                users[user_id]["is_authenticated"] = True
                users[user_id]["failed_attempts"] = 0
            
            save_users(users)
            log_event(user_id, "login_success")
            
            await update.message.reply_text("ورود موفقیت‌آمیز!")
            await show_menu(update, context)
        else:
            # پس‌کی نامعتبر است
            if user_id in users:
                users[user_id]["failed_attempts"] = users[user_id].get("failed_attempts", 0) + 1
                save_users(users)
            
            log_event(user_id, "login_failed")
            
            await update.message.reply_text(
                "پس‌کی نامعتبر است. لطفاً دوباره تلاش کنید:"
            )
    
    elif state == WAITING_FOR_REPORT:
        # دریافت گزارش از کاربر
        try:
            # افزودن پیام راهنما به متن ارسالی کاربر برای ChatGPT
            prompt = f"""
            لطفاً متن زیر را به یک گزارش یا نامه رسمی برای شورای صنفی دانشگاه اصفهان تبدیل کنید. 
            متن را کامل، مرتب و با رعایت اصول نگارش رسمی بازنویسی کنید.
            
            متن کاربر:
            {text}
            """
            
            # ارسال به ChatGPT
            ai_response = ai_config.ask_ai(prompt)
            
            # ذخیره گزارش
            report_id = save_report(user_id, text, ai_response)
            
            # ذخیره گزارش در context برای استفاده بعدی
            context.user_data["current_report"] = ai_response
            context.user_data["report_id"] = report_id
            
            # ایجاد دکمه‌های تأیید
            keyboard = [
                [InlineKeyboardButton("تایید پیش‌نویس", callback_data="confirm")],
                [InlineKeyboardButton("ویرایش متن", callback_data="edit")]
            ]
            
            # نمایش پیش‌نویس به کاربر
            await update.message.reply_text(
                f"پیش‌نویس شما آماده شد:\n\n{ai_response}",
                reply_markup=InlineKeyboardMarkup(keyboard)
            )
            
            # تغییر حالت کاربر
            context.user_data["state"] = WAITING_FOR_CONFIRMATION
            log_event(user_id, "draft_generated")
            
        except Exception as e:
            # خطا در ارتباط با ChatGPT
            logger.error(f"Error generating AI response: {e}")
            log_event(user_id, "ai_error", str(e))
            await update.message.reply_text(
                "متأسفانه در پردازش درخواست شما خطایی رخ داده است. لطفاً دوباره تلاش کنید."
            )
            await show_menu(update, context)
    
    elif state == NORMAL_MODE:
        # پردازش گزینه‌های منو
        if text == "ارسال گزارش/نامه":
            context.user_data["state"] = WAITING_FOR_REPORT
            await update.message.reply_text(
                "لطفاً متن اولیه گزارش یا نامه خود را وارد کنید:"
            )
        elif text == "راهنما":
            await show_help(update, context)
        elif text == "آرشیو":
            await show_archive(update, context)
        elif text == "ارتباط با مدیر":
            await contact_admin(update, context)
        else:
            await update.message.reply_text(
                "دستور نامعتبر. لطفاً از منوی موجود استفاده کنید."
            )

# پردازش دکمه‌های اینلاین
async def handle_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = str(update.effective_user.id)
    query = update.callback_query
    await query.answer()
    
    if query.data == "confirm":
        # تأیید پیش‌نویس
        report_content = context.user_data.get("current_report", "")
        report_id = context.user_data.get("report_id", 0)
        
        # ایجاد فایل ورد
        docx_file = create_docx(report_content, report_id)
        
        # ارسال فایل به کاربر
        await query.message.reply_document(
            document=docx_file,
            filename=f"report_{report_id}.docx",
            caption="فایل گزارش شما آماده شد."
        )
        
        log_event(user_id, "report_confirmed", {"report_id": report_id})
        await show_menu(update, context)
    
    elif query.data == "edit":
        # ویرایش متن
        context.user_data["state"] = WAITING_FOR_REPORT
        await query.message.reply_text(
            "لطفاً متن جدید خود را وارد کنید:"
        )

# نمایش راهنما
async def show_help(update: Update, context: ContextTypes.DEFAULT_TYPE):
    help_text = """🔰 راهنمای استفاده از ربات شورای صنفی 🔰

- برای ارسال گزارش یا نامه، گزینه "ارسال گزارش/نامه" را انتخاب کنید.
- پس از وارد کردن متن اولیه، پیش‌نویس توسط هوش مصنوعی تکمیل می‌شود.
- در صورت تایید پیش‌نویس، فایل Word نهایی برای شما ارسال می‌شود.
- برای مشاهده گزارش‌های قبلی، گزینه "آرشیو" را انتخاب کنید.
- برای ارتباط با مدیر، گزینه "ارتباط با مدیر" را انتخاب کنید.

🔑 توجه: لطفاً پس‌کی خود را در اختیار دیگران قرار ندهید.
"""
    await update.message.reply_text(help_text)

# نمایش آرشیو
async def show_archive(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = str(update.effective_user.id)
    
    try:
        with open(REPORTS_FILE, 'r', encoding='utf-8') as f:
            reports = json.load(f)
        
        # فیلتر کردن گزارش‌های کاربر فعلی
        user_reports = [r for r in reports if r["user_id"] == user_id]
        
        if not user_reports:
            await update.message.reply_text("شما هنوز گزارشی ثبت نکرده‌اید.")
            return
        
        # نمایش ۵ گزارش آخر
        recent_reports = user_reports[-5:]
        response = "گزارش‌های اخیر شما:\n\n"
        
        for report in recent_reports:
            date = datetime.fromisoformat(report["timestamp"]).strftime("%Y/%m/%d")
            report_brief = report["ai_response"][:50] + "..." if len(report["ai_response"]) > 50 else report["ai_response"]
            response += f"📄 شناسه گزارش: {report['id']}\n"
            response += f"📅 تاریخ: {date}\n"
            response += f"📝 خلاصه: {report_brief}\n\n"
        
        # ایجاد دکمه‌های آرشیو
        buttons = []
        for report in recent_reports:
            buttons.append([InlineKeyboardButton(f"دریافت گزارش {report['id']}", callback_data=f"get_report_{report['id']}")])
        
        await update.message.reply_text(
            response,
            reply_markup=InlineKeyboardMarkup(buttons)
        )
    
    except Exception as e:
        logger.error(f"Error showing archive: {e}")
        await update.message.reply_text("خطا در نمایش آرشیو. لطفاً دوباره تلاش کنید.")

# ارتباط با مدیر
async def contact_admin(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text(
        "برای ارتباط با مدیر، لطفاً پیام خود را با /admin شروع کنید. مثال:\n"
        "/admin سلام، یک سوال داشتم..."
    )

# ارسال پیام به مدیر
async def send_to_admin(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.effective_user.id
    user_name = update.effective_user.first_name
    message = update.message.text[7:]  # حذف /admin از ابتدای پیام
    
    # ارسال پیام به ادمین
    try:
        await context.bot.send_message(
            chat_id=ADMIN_USER_ID,
            text=f"پیام جدید از کاربر {user_name} (ID: {user_id}):\n\n{message}"
        )
        await update.message.reply_text("پیام شما با موفقیت به مدیر ارسال شد.")
    except Exception as e:
        logger.error(f"Error sending message to admin: {e}")
        await update.message.reply_text("خطا در ارسال پیام به مدیر. لطفاً دوباره تلاش کنید.")

# پاسخ مدیر به کاربر
async def admin_reply(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.effective_user.id
    
    # بررسی این که آیا فرستنده مدیر است یا خیر
    if user_id != ADMIN_USER_ID:
        await update.message.reply_text("شما دسترسی به این دستور را ندارید.")
        return
    
    # فرمت دستور: /reply ID MESSAGE
    try:
        command_parts = update.message.text.split(" ", 2)
        target_user_id = command_parts[1]
        message = command_parts[2]
        
        await context.bot.send_message(
            chat_id=target_user_id,
            text=f"پاسخ از مدیر:\n\n{message}"
        )
        await update.message.reply_text(f"پاسخ شما به کاربر {target_user_id} ارسال شد.")
    except Exception as e:
        logger.error(f"Error in admin reply: {e}")
        await update.message.reply_text("خطا در ارسال پاسخ. لطفاً دستور را به شکل صحیح وارد کنید:\n/reply ID MESSAGE")

# دستور آمار (فقط برای مدیر)
async def stats(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.effective_user.id
    
    # بررسی این که آیا فرستنده مدیر است یا خیر
    if user_id != ADMIN_USER_ID:
        await update.message.reply_text("شما دسترسی به این دستور را ندارید.")
        return
    
    try:
        # آمار کاربران
        users = load_users()
        active_users = sum(1 for u in users.values() if u.get("is_authenticated", False))
        
        # آمار گزارش‌ها
        try:
            with open(REPORTS_FILE, 'r', encoding='utf-8') as f:
                reports = json.load(f)
            total_reports = len(reports)
        except (FileNotFoundError, json.JSONDecodeError):
            total_reports = 0
        
        # آمار لاگ‌ها
        try:
            with open(LOGS_FILE, 'r', encoding='utf-8') as f:
                logs = json.load(f)
            total_logs = len(logs)
            login_failures = sum(1 for log in logs if log["event_type"] == "login_failed")
            ai_errors = sum(1 for log in logs if log["event_type"] == "ai_error")
        except (FileNotFoundError, json.JSONDecodeError):
            total_logs = 0
            login_failures = 0
            ai_errors = 0
        
        stats_text = f"""📊 آمار سیستم:

👥 تعداد کل کاربران: {len(users)}
✅ کاربران فعال: {active_users}

📝 گزارش‌های ثبت شده: {total_reports}

🔑 تلاش‌های ناموفق ورود: {login_failures}
⚠️ خطاهای هوش مصنوعی: {ai_errors}
📜 کل رویدادهای ثبت شده: {total_logs}
"""
        await update.message.reply_text(stats_text)
    
    except Exception as e:
        logger.error(f"Error generating stats: {e}")
        await update.message.reply_text("خطا در تولید آمار. لطفاً دوباره تلاش کنید.")

def main():
    # ایجاد و پیکربندی اپلیکیشن
    application = Application.builder().token(TOKEN).build()
    
    # اضافه کردن هندلرها
    application.add_handler(CommandHandler("start", start))
    application.add_handler(CommandHandler("help", show_help))
    application.add_handler(CommandHandler("admin", send_to_admin))
    application.add_handler(CommandHandler("reply", admin_reply))
    application.add_handler(CommandHandler("stats", stats))
    application.add_handler(CallbackQueryHandler(handle_callback))
    application.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_message))
    
    # شروع بات
    application.run_polling()

if __name__ == "__main__":
    main()